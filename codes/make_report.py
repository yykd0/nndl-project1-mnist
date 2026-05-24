import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SUBMISSION_DIR = BASE_DIR.parent.parent / "submission"
FIGURE_DIR = SUBMISSION_DIR / "figures"
RESULTS_PATH = SUBMISSION_DIR / "experiment_results.json"
REPORT_PATH = SUBMISSION_DIR / "PJ1_report.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def set_paragraph_font(paragraph, size=11):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_text(cell, text, bold=False, fill=None):
    text = str(text)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = "#2E74B5" if level <= 2 else "#1F4D78"
    size = 16 if level == 1 else 13 if level == 2 else 12
    for run in p.runs:
        set_run_font(run, size=size, bold=True, color=color)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_font(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, color="#555555")
    return p


def add_picture_if_exists(doc, path, caption, width=5.9):
    if Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def pct(value):
    return f"{value * 100:.2f}%"


def seconds(value):
    if value < 60:
        return f"{value:.1f} 秒"
    return f"{value / 60:.1f} 分钟"


def result_map(results):
    return {item["name"]: item for item in results}


def add_results_table(doc, results):
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["实验", "模型", "优化器", "轮数", "训练准确率", "验证准确率", "测试准确率", "耗时"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, fill="#F2F4F7")

    model_names = {
        "mlp_sgd": "MLP",
        "cnn_sgd": "CNN",
        "cnn_momentum": "CNN",
    }
    optimizer_names = {
        "sgd": "SGD",
        "momentum": "MomentumGD",
    }
    scheduler_names = {
        "none": "无",
        "multistep": "MultiStepLR",
        "exponential": "ExponentialLR",
    }
    for item in results:
        row = table.add_row().cells
        optimizer = optimizer_names.get(item["optimizer"], item["optimizer"])
        if item["scheduler"] != "none":
            optimizer = f"{optimizer} + {scheduler_names.get(item['scheduler'], item['scheduler'])}"
        values = [
            item["name"],
            model_names.get(item["name"], item["name"]),
            optimizer,
            item["epochs"],
            pct(item["train_subset_accuracy"]),
            pct(item["valid_accuracy"]),
            pct(item["test_accuracy"]),
            seconds(item["train_time_seconds"]),
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "#2E74B5"),
        ("Heading 2", 13, "#2E74B5"),
        ("Heading 3", 12, "#1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.font.bold = True


def build_report():
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        payload = json.load(f)
    results = payload["results"]
    figures = payload.get("figures", {})
    by_name = result_map(results)

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project 1：基于 MLP 与 CNN 的 MNIST 手写数字分类")
    set_run_font(run, size=22, bold=True, color="#0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("神经网络与深度学习课程作业")
    set_run_font(run, size=12, color="#555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("姓名：[姓名待补充]    学号：[学号待补充]")
    set_run_font(run, size=11)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run("Github 代码链接：https://github.com/yykd0/nndl-project1-mnist    模型权重链接：[ModelScope/checkpoint 链接待补充]")
    set_run_font(run, size=10, color="#555555")

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本项目使用课程提供的 MNIST 数据集，基于 NumPy 从基础组件开始实现手写数字分类模型。"
        "实验包括一个 MLP 基线模型、一个自实现卷积算子的 CNN 模型，以及两个附加方向：优化方法比较和错误分析可视化。"
        "实验过程中没有调用 PyTorch 等深度学习框架的现成算子，也没有使用外部数据集。"
    )
    if payload.get("quick"):
        add_body(doc, "注意：当前结果由 quick 模式生成，仅用于流程检查；正式提交前应运行完整实验。")

    add_heading(doc, "1. MLP 基线模型", 1)
    add_body(
        doc,
        "MLP 基线采用 784-256-10 的两层全连接结构。输入为拉平成 784 维向量的 28×28 灰度图像，"
        "隐藏层使用 ReLU 激活，输出层给出 10 个类别的 logits。Linear 层的前向传播、反向传播、"
        "权重梯度和偏置梯度均使用 NumPy 手动实现。"
    )
    add_bullet(doc, "损失函数：数值稳定的 softmax 多分类交叉熵。")
    add_bullet(doc, "数据划分：固定随机种子，将训练集划分为 50,000 张训练图像和 10,000 张验证图像。")
    add_bullet(doc, "优化器：mini-batch SGD。")
    mlp = by_name.get("mlp_sgd")
    if mlp:
        add_body(
            doc,
            f"MLP 在验证集上的准确率为 {pct(mlp['valid_accuracy'])}，测试集准确率为 {pct(mlp['test_accuracy'])}，"
            f"训练子集准确率为 {pct(mlp['train_subset_accuracy'])}。"
        )

    add_heading(doc, "2. CNN 模型与 MLP 对比", 1)
    add_body(
        doc,
        "CNN 直接利用图像的二维空间结构，模型结构为 Conv2D(1->8, 5×5) -> ReLU -> 2×2 MaxPool -> "
        "Flatten -> Linear(64 hidden units) -> ReLU -> Linear(10 classes)。卷积层采用 im2col 风格的 "
        "NumPy 向量化实现，反向传播中对输入梯度进行 scatter-add 累加，从而兼顾实现清晰度和训练效率。"
    )
    cnn = by_name.get("cnn_sgd")
    if cnn and mlp:
        diff = cnn["test_accuracy"] - mlp["test_accuracy"]
        add_body(
            doc,
            f"在相近训练设置下，CNN-SGD 的测试准确率为 {pct(cnn['test_accuracy'])}，"
            f"相对 MLP 基线提升 {diff * 100:+.2f} 个百分点。CNN 更适合图像分类的主要原因是卷积核能够提取局部笔画模式，"
            "并在不同空间位置共享参数；而 MLP 将图像完全拉平，较难显式保留相邻像素之间的局部结构。"
        )

    add_heading(doc, "3. 附加方向一：优化方法", 1)
    opt = by_name.get("cnn_momentum")
    if opt and cnn:
        diff = opt["test_accuracy"] - cnn["test_accuracy"]
        add_body(
            doc,
            "优化方向比较普通 SGD 与 MomentumGD + MultiStepLR 在同一 CNN 结构上的表现。Momentum 可以累积历史梯度方向，"
            "减小 mini-batch 随机噪声带来的震荡；学习率调度则在训练后期降低步长，使模型更稳定地收敛。"
        )
        add_body(
            doc,
            f"加入 Momentum 与学习率调度后，CNN 在验证集上的准确率为 {pct(opt['valid_accuracy'])}，"
            f"测试集准确率为 {pct(opt['test_accuracy'])}，相比 CNN-SGD 测试准确率变化为 {diff * 100:+.2f} 个百分点。"
        )

    add_heading(doc, "4. 附加方向二：错误分析与可视化", 1)
    add_body(
        doc,
        "仅报告一个最终准确率并不能充分解释模型行为，因此本项目进一步绘制学习曲线、混淆矩阵、误分类样例、"
        "CNN 卷积核和 MLP 首层权重。学习曲线用于观察收敛过程，混淆矩阵展示类别间的系统性混淆，"
        "误分类样例则帮助分析哪些笔画形态仍然困难。"
    )
    add_picture_if_exists(doc, figures.get("validation_accuracy", FIGURE_DIR / "validation_accuracy.png"), "图 1  验证集准确率随训练过程变化。")
    add_picture_if_exists(doc, figures.get("training_loss", FIGURE_DIR / "training_loss.png"), "图 2  训练损失曲线。")
    add_picture_if_exists(doc, figures.get("confusion_matrix", FIGURE_DIR / "confusion_matrix.png"), "图 3  CNN + Momentum 在测试集上的混淆矩阵。")
    add_picture_if_exists(doc, figures.get("misclassified_examples", FIGURE_DIR / "misclassified_examples.png"), "图 4  部分误分类测试样例。")
    add_picture_if_exists(doc, figures.get("cnn_kernels", FIGURE_DIR / "cnn_kernels.png"), "图 5  CNN 第一层卷积核可视化。", width=5.2)
    add_picture_if_exists(doc, figures.get("mlp_first_layer_weights", FIGURE_DIR / "mlp_first_layer_weights.png"), "图 6  MLP 第一层权重可视化。", width=5.2)

    add_heading(doc, "5. 主要实验结果", 1)
    add_results_table(doc, results)

    add_heading(doc, "6. 讨论", 1)
    add_body(
        doc,
        "从结果看，MLP 已经可以在 MNIST 上取得较高准确率，这是因为 MNIST 图像较干净、居中且类别区分度较高。"
        "但 CNN 在相近训练预算下进一步提升了表现，说明局部感受野和参数共享对图像任务是有效归纳偏置。"
        "在 CNN 上加入 Momentum 和学习率调度后，验证集和测试集准确率进一步提高，表明优化策略对收敛速度和最终性能都有帮助。"
    )
    add_body(
        doc,
        "从错误分析看，模型较容易混淆笔画形态相近或书写不规范的数字。例如某些 9、4、7、8 的写法存在局部结构相似，"
        "部分样本还存在倾斜、断笔或笔画过粗/过细的问题。混淆矩阵可以给出整体错误分布，误分类样例则能直观说明这些错误的视觉原因。"
    )

    add_heading(doc, "复现与提交说明", 1)
    add_bullet(doc, "运行 `sanity_checks.py` 可检查 Linear 与 Conv2D 的梯度实现。")
    add_bullet(doc, "运行 `test_train.py` 可重新生成 checkpoint、实验指标和图表。")
    add_bullet(doc, "上传 Github 时只上传代码包，不要上传 MNIST 数据集、模型权重、生成图表或其他大文件。")
    add_bullet(doc, "将 `checkpoints/` 中的模型文件上传到 ModelScope 或其他权重托管平台后，替换报告首页的模型链接占位。")

    doc.save(REPORT_PATH)
    print(f"saved={REPORT_PATH}")


if __name__ == "__main__":
    build_report()
